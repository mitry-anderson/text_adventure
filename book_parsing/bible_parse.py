import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

folder = "kjv"
filename = "mark"

run = True
chapter_num = 1
chapter_regex = r"\d*(?=:)"
chapter_and_verse_regex = r"[\d*:\d*]"

output_doc = Document()


title = output_doc.add_paragraph()
title_run = title.add_run("\n" + filename.upper())
title_font = title_run.font
title_font.name = "Baskerville Old Face"
title_font.size = Pt(90)
title_font.bold = True
title_font.underline = False
title_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def write_paragraph(output_line):
    f_out.write(output_line + "\n")
    para = output_doc.add_paragraph().add_run(output_line)
    para.font.name = "Baskerville Old Face"
    para.font.size = Pt(14)
    print(output_line)

def write_chapter(num):
    output_doc.add_page_break()
    output_line = "\nCHAPTER " + str(num)
    f_out.write(output_line + "\n")
    head = output_doc.add_paragraph()
    head.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head_run = head.add_run(output_line)
    head_font = head_run.font
    head_font.name = "Baskerville Old Face"
    head_font.size = Pt(20)
    head_font.bold = True
    head_font.underline = False
    head_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    print(output_line)

with open(folder + "/" + filename + ".txt", "r") as f_in:
    with open(folder + "/" + "edited_" + filename + ".txt", "w") as f_out:
        write_chapter(chapter_num)
        output_line = ""
        while(run):
            line = f_in.readline()

            if (line == ""):
                # stop and write out the last paragraph
                run = False
                write_paragraph(output_line)
            else:
                # see if this line is a chapter heading
                chap = re.findall(chapter_regex, line)
                chapter_num_curr = int(chap[0]) if len(chap) > 0 else chapter_num
                if(chapter_num_curr > chapter_num):
                    # write paragraph + next chapter
                    write_paragraph(output_line)
                    chapter_num = chapter_num_curr
                    write_chapter(chapter_num)
                    output_line = ""

                # regex messarounds
                edited_line = re.sub(chapter_and_verse_regex,'',line)
                edited_line = re.sub("\n", "", edited_line)
                if edited_line != "":
                    output_line += edited_line + "\n"

# save the word doc too bruh
output_doc.save(folder + '/' + filename + '.docx')













